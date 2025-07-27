from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime
from typing import Dict, List, Optional
import config

class DocumentGenerator:
    """Generates Word documents with formal logic arguments and critiques."""
    
    def __init__(self):
        """Initialize the document generator."""
        pass
    
    def generate_documents(self, logic_data: Dict, include_critiques: bool = False, 
                          output_filename: str = None) -> List[str]:
        """
        Generate Word documents with formal logic arguments.
        
        Args:
            logic_data: Dictionary containing formal logic arguments and critiques
            include_critiques: Whether to include critiques (generates two documents if True)
            output_filename: Base filename for output documents
            
        Returns:
            List of generated document file paths
        """
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"debate_analysis_{timestamp}"
        
        generated_files = []
        
        # Generate clean document (without critiques)
        clean_doc_path = self._generate_clean_document(logic_data, output_filename)
        generated_files.append(clean_doc_path)
        
        # Generate critiqued document if requested
        if include_critiques and "critiques" in logic_data:
            critique_doc_path = self._generate_critique_document(logic_data, output_filename)
            generated_files.append(critique_doc_path)
        
        return generated_files
    
    def _generate_clean_document(self, logic_data: Dict, base_filename: str) -> str:
        """Generate a clean document without critiques."""
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Add title
        title = doc.add_heading('Formal Logic Analysis of Debate', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        self._add_metadata_section(doc, logic_data)
        
        # Add formal arguments by speaker
        self._add_speaker_arguments(doc, logic_data["formal_arguments"], include_critiques=False)
        
        # Add original transcription as appendix
        self._add_transcription_appendix(doc, logic_data.get("original_transcription", {}))
        
        # Save document
        filename = f"{base_filename}_clean.docx"
        filepath = os.path.join(config.OUTPUT_DIR, filename)
        doc.save(filepath)
        
        return filepath
    
    def _generate_critique_document(self, logic_data: Dict, base_filename: str) -> str:
        """Generate a document with critiques and highlighted problems."""
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Add title
        title = doc.add_heading('Formal Logic Analysis of Debate (With Critiques)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        self._add_metadata_section(doc, logic_data)
        
        # Add formal arguments by speaker with critiques
        self._add_speaker_arguments(doc, logic_data["formal_arguments"], 
                                  include_critiques=True, critiques=logic_data.get("critiques", {}))
        
        # Add critique summary
        self._add_critique_summary(doc, logic_data.get("critiques", {}))
        
        # Add original transcription as appendix
        self._add_transcription_appendix(doc, logic_data.get("original_transcription", {}))
        
        # Save document
        filename = f"{base_filename}_with_critiques.docx"
        filepath = os.path.join(config.OUTPUT_DIR, filename)
        doc.save(filepath)
        
        return filepath
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for the document."""
        # Create custom styles
        styles = doc.styles
        
        # Problem highlight style
        try:
            problem_style = styles.add_style('ProblemHighlight', WD_STYLE_TYPE.CHARACTER)
            problem_style.font.color.rgb = RGBColor(255, 0, 0)  # Red text
            problem_style.font.bold = True
        except:
            pass  # Style might already exist
        
        # Critique note style
        try:
            critique_style = styles.add_style('CritiqueNote', WD_STYLE_TYPE.PARAGRAPH)
            critique_style.font.color.rgb = RGBColor(128, 0, 0)  # Dark red
            critique_style.font.italic = True
            critique_style.font.size = 10
        except:
            pass
    
    def _add_metadata_section(self, doc: Document, logic_data: Dict):
        """Add metadata section to the document."""
        doc.add_heading('Document Information', level=1)
        
        # Add generation timestamp
        p = doc.add_paragraph()
        p.add_run('Generated: ').bold = True
        p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        
        # Add speaker information
        speakers = logic_data.get("speakers", [])
        if speakers:
            p = doc.add_paragraph()
            p.add_run('Speakers: ').bold = True
            p.add_run(", ".join(speakers))
        
        # Add transcription language if available
        orig_trans = logic_data.get("original_transcription", {})
        if "language" in orig_trans:
            p = doc.add_paragraph()
            p.add_run('Detected Language: ').bold = True
            p.add_run(orig_trans["language"])
        
        doc.add_page_break()
    
    def _add_speaker_arguments(self, doc: Document, formal_arguments: Dict, 
                             include_critiques: bool = False, critiques: Dict = None):
        """Add formal arguments section organized by speaker."""
        doc.add_heading('Formal Logic Arguments by Speaker', level=1)
        
        for speaker, arguments in formal_arguments.items():
            # Speaker heading
            doc.add_heading(f'{speaker}', level=2)
            
            # Add speaker's original text
            if "original_text" in arguments:
                doc.add_heading('Original Statement:', level=3)
                p = doc.add_paragraph(arguments["original_text"])
                p.style = 'Italic'
            
            # Add formal analysis
            if "raw_analysis" in arguments:
                doc.add_heading('Formal Logic Analysis:', level=3)
                
                analysis_text = arguments["raw_analysis"]
                
                if include_critiques and critiques and speaker in critiques:
                    # Highlight problematic sections
                    self._add_highlighted_analysis(doc, analysis_text, critiques[speaker])
                else:
                    # Add clean analysis
                    doc.add_paragraph(analysis_text)
            
            # Add structured arguments if available
            if "structured_arguments" in arguments and arguments["structured_arguments"]:
                self._add_structured_arguments(doc, arguments["structured_arguments"])
            
            # Add critiques if requested
            if include_critiques and critiques and speaker in critiques:
                self._add_speaker_critique(doc, critiques[speaker])
            
            doc.add_page_break()
    
    def _add_highlighted_analysis(self, doc: Document, analysis_text: str, critique_data: Dict):
        """Add analysis text with problematic sections highlighted in red."""
        # For simplicity, we'll add the full analysis and then add critique notes
        # In a more sophisticated implementation, you'd parse and highlight specific sentences
        
        doc.add_paragraph(analysis_text)
        
        # Add critique problems as highlighted notes
        problems = critique_data.get("identified_problems", [])
        if problems:
            doc.add_heading('⚠️ Logical Issues Identified:', level=4)
            
            for i, problem in enumerate(problems, 1):
                p = doc.add_paragraph()
                
                # Problem number and type
                run = p.add_run(f"Problem {i}: {problem.get('problem_type', 'Issue')}")
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                run.bold = True
                
                # Problem description
                p.add_run(f"\n{problem.get('description', 'No description available')}")
                
                # Severity
                severity = problem.get('severity', 'Unknown')
                run = p.add_run(f"\nSeverity: {severity}")
                run.italic = True
    
    def _add_structured_arguments(self, doc: Document, structured_args: List[Dict]):
        """Add structured argument breakdown."""
        doc.add_heading('Structured Argument Breakdown:', level=3)
        
        for i, arg in enumerate(structured_args, 1):
            doc.add_heading(f'Argument {i}', level=4)
            
            # Premises
            if arg.get("premises"):
                p = doc.add_paragraph()
                p.add_run('Premises:').bold = True
                for premise in arg["premises"]:
                    doc.add_paragraph(f"• {premise}", style='List Bullet')
            
            # Conclusions
            if arg.get("conclusions"):
                p = doc.add_paragraph()
                p.add_run('Conclusions:').bold = True
                for conclusion in arg["conclusions"]:
                    doc.add_paragraph(f"• {conclusion}", style='List Bullet')
            
            # Logical structure
            if arg.get("logical_structure"):
                p = doc.add_paragraph()
                p.add_run('Logical Structure: ').bold = True
                p.add_run(arg["logical_structure"])
            
            # Argument type
            if arg.get("argument_type"):
                p = doc.add_paragraph()
                p.add_run('Argument Type: ').bold = True
                p.add_run(arg["argument_type"])
            
            # Supporting evidence
            if arg.get("supporting_evidence"):
                p = doc.add_paragraph()
                p.add_run('Supporting Evidence: ').bold = True
                p.add_run(arg["supporting_evidence"])
    
    def _add_speaker_critique(self, doc: Document, critique_data: Dict):
        """Add critique section for a speaker."""
        doc.add_heading('Critical Analysis:', level=3)
        
        if "raw_critique" in critique_data:
            doc.add_paragraph(critique_data["raw_critique"])
    
    def _add_critique_summary(self, doc: Document, critiques: Dict):
        """Add overall critique summary."""
        doc.add_heading('Overall Critique Summary', level=1)
        
        total_problems = 0
        severity_counts = {"Critical": 0, "Major": 0, "Moderate": 0, "Minor": 0}
        
        # Count problems across all speakers
        for speaker, critique_data in critiques.items():
            problems = critique_data.get("identified_problems", [])
            total_problems += len(problems)
            
            for problem in problems:
                severity = problem.get("severity", "Unknown")
                if severity in severity_counts:
                    severity_counts[severity] += 1
        
        # Add summary statistics
        p = doc.add_paragraph()
        p.add_run('Total Logical Issues Identified: ').bold = True
        p.add_run(str(total_problems))
        
        # Severity breakdown
        doc.add_heading('Issues by Severity:', level=2)
        for severity, count in severity_counts.items():
            if count > 0:
                p = doc.add_paragraph(f"{severity}: {count}", style='List Bullet')
        
        # Speaker breakdown
        doc.add_heading('Issues by Speaker:', level=2)
        for speaker, critique_data in critiques.items():
            problems = critique_data.get("identified_problems", [])
            p = doc.add_paragraph(f"{speaker}: {len(problems)} issues", style='List Bullet')
    
    def _add_transcription_appendix(self, doc: Document, transcription_data: Dict):
        """Add original transcription as appendix."""
        doc.add_heading('Appendix: Original Transcription', level=1)
        
        segments = transcription_data.get("segments", [])
        if not segments:
            doc.add_paragraph("No transcription data available.")
            return
        
        # Add full text
        doc.add_heading('Complete Transcript:', level=2)
        full_text = transcription_data.get("text", "")
        if full_text:
            doc.add_paragraph(full_text)
        
        # Add segment breakdown
        doc.add_heading('Segment Breakdown:', level=2)
        
        current_speaker = None
        for segment in segments:
            speaker = segment.get("speaker", "Unknown")
            start_time = segment.get("start", 0)
            end_time = segment.get("end", 0)
            text = segment.get("text", "")
            
            # Add speaker change
            if speaker != current_speaker:
                p = doc.add_paragraph()
                p.add_run(f"\n{speaker}:").bold = True
                current_speaker = speaker
            
            # Add timestamped text
            p = doc.add_paragraph()
            p.add_run(f"[{start_time:.1f}s - {end_time:.1f}s] ").italic = True
            p.add_run(text)
    
    def create_sample_document(self, output_path: str = None) -> str:
        """Create a sample document for testing purposes."""
        if output_path is None:
            output_path = os.path.join(config.OUTPUT_DIR, "sample_analysis.docx")
        
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add title
        title = doc.add_heading('Sample Formal Logic Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sample content
        doc.add_heading('Sample Speaker Analysis', level=1)
        
        doc.add_heading('Speaker 1', level=2)
        doc.add_paragraph("Sample argument: All men are mortal. Socrates is a man. Therefore, Socrates is mortal.")
        
        doc.add_heading('Formal Analysis:', level=3)
        doc.add_paragraph("This is a classic example of a valid deductive syllogism.")
        
        doc.save(output_path)
        return output_path 